from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


class DOCXParser:
    def parse(self, file_path: str) -> dict:
        source = Path(file_path)
        if not source.is_file():
            raise FileNotFoundError(f"Файл DOCX не найден: {file_path}")
        document = Document(source)
        parts: list[str] = []
        links: list[str] = []
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(f"Заголовок документа: {text}")
                links.extend(self._paragraph_links(paragraph))
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                style_name = paragraph.style.name.lower() if paragraph.style else ""
                prefix = "- " if "list" in style_name or "спис" in style_name else ""
                parts.append(f"{prefix}{text}")
            links.extend(self._paragraph_links(paragraph))
        tables: list[dict] = []
        for table_number, table in enumerate(document.tables, start=1):
            values = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not values:
                continue
            headers = [value or f"column_{index + 1}" for index, value in enumerate(values[0])]
            rows = [dict(zip(headers, row)) for row in values[1:]]
            tables.append({"table": table_number, "headers": headers, "rows": rows})
            parts.append("\n".join(" | ".join(row) for row in values))
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        links.extend(self._paragraph_links(paragraph))
        return {"raw_text": "\n\n".join(parts), "tables": tables, "links": links, "excel_audit": None, "image_count": len(document.inline_shapes)}

    @staticmethod
    def _paragraph_links(paragraph) -> list[str]:
        links: list[str] = []
        for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
            relation_id = hyperlink.get(qn("r:id"))
            if relation_id and relation_id in paragraph.part.rels:
                target = paragraph.part.rels[relation_id].target_ref
                if target:
                    links.append(target)
        return links
