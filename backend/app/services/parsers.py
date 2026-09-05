import re
from pathlib import Path

import docx
import openpyxl
import pdfplumber

from .contracts import Constraints, Criterion, TaskRubric


SUPPORTED_TASK_EXTENSIONS = frozenset({".pdf", ".docx", ".xlsx", ".md"})


def extract_task_text(path: str | Path) -> str:
    source = Path(path)
    extension = source.suffix.lower()
    if not source.is_file():
        raise ValueError("Task file does not exist")
    if extension == ".pdf":
        with pdfplumber.open(source) as pdf:
            return "\n\n".join((page.extract_text() or "") for page in pdf.pages).strip()
    if extension == ".docx":
        document = docx.Document(source)
        text = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            text.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(text).strip()
    if extension == ".xlsx":
        workbook = openpyxl.load_workbook(source, data_only=False)
        rows = []
        for sheet in workbook.worksheets:
            rows.append(f"# {sheet.title}")
            rows.extend(" | ".join(str(value or "") for value in row) for row in sheet.iter_rows(values_only=True))
        return "\n".join(rows).strip()
    if extension == ".md":
        return source.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported task file extension: {extension}")


def fallback_rubric(task_id: str, title: str, text: str, criteria: list[dict] | None = None) -> TaskRubric:
    parsed = [
        Criterion(
            name=str(item.get("title", item.get("name", "Criterion"))),
            description=str(item.get("description", "")),
            max_points=float(item.get("max_score", item.get("max_points", 0))),
        )
        for item in criteria or []
    ]
    if not parsed:
        parsed = _extract_criteria(text)
    return TaskRubric(
        task_id=task_id,
        title=title,
        description=_first_paragraph(text),
        full_instructions=text,
        guidelines=_extract_guidelines(text),
        criteria=parsed,
        constraints=Constraints(),
        total_points=sum(item.max_points for item in parsed),
    )


def _extract_criteria(text: str) -> list[Criterion]:
    result = []
    pattern = re.compile(r"(?:^|\n)\s*(?:\d+[.)]|[-*])\s*([^\n]+?)(?:\s*[-—:]\s*)(\d+(?:[.,]\d+)?)\s*(?:балл|point)", re.I)
    for match in pattern.finditer(text):
        result.append(Criterion(name=match.group(1).strip(), max_points=float(match.group(2).replace(",", "."))))
    return result


def _extract_guidelines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if re.match(r"^\s*\d+[.)]", line)]


def _first_paragraph(text: str) -> str:
    return next((line.strip() for line in text.splitlines() if line.strip()), "")
